
"""
Enhanced RAG Data Loader with support for PDF, TXT, and DOCX files.
Optimized for ChromaDB integration with consistent collection management.
"""

import os
import re
import logging
import time
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from enum import Enum

import openai
import chromadb
from dotenv import load_dotenv
import PyPDF2
import tiktoken

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
data_dir = os.getenv("DATA_DIR", "data")

class ChunkingStrategy(Enum):
    """Available text chunking strategies."""
    SENTENCE = "sentence"
    RECURSIVE = "recursive" 
    FIXED_SIZE = "fixed_size"

class FileType(Enum):
    """Supported file types."""
    TXT = "txt"
    PDF = "pdf"
    DOCX = "docx"

@dataclass
class ChunkingConfig:
    """Configuration for text chunking."""
    strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE
    chunk_size: int = 1000
    chunk_overlap: int = 200
    separators: Optional[List[str]] = None
    
    def __post_init__(self):
        if self.separators is None:
            self.separators = ["\n\n", "\n", ". ", " ", ""]

@dataclass
class EmbeddingConfig:
    """Configuration for embeddings."""
    model: str = "text-embedding-3-small"
    batch_size: int = 50
    max_tokens: int = 8191

class DocumentProcessor:
    """Handles document processing and text extraction."""
    
    @staticmethod
    def extract_text_from_txt(file_path: Path) -> str:
        """Extract text from a TXT file."""
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                try:
                    return file_path.read_text(encoding=encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file {file_path}")
    
    @staticmethod
    def extract_text_from_pdf(file_path: Path) -> str:
        """Extract text from PDF using PyPDF2."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1} from {file_path}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            raise
        
        if not text.strip():
            raise ValueError(f"No text could be extracted from PDF: {file_path}")
        
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            import docx
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            if not full_text.strip():
                raise ValueError(f"No text could be extracted from DOCX: {file_path}")
            
            return full_text
            
        except ImportError:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
    
    @classmethod
    def extract_text(cls, file_path: Path) -> str:
        """Extract text from supported file formats."""
        extension = file_path.suffix.lower().lstrip('.')
        
        if extension == 'txt':
            return cls.extract_text_from_txt(file_path)
        elif extension == 'pdf':
            return cls.extract_text_from_pdf(file_path)
        elif extension == 'docx':
            return cls.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

class TextChunker:
    """Handles text chunking with different strategies."""
    
    def __init__(self, config: ChunkingConfig):
        self.config = config
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.encoding.encode(text))
    
    def _clean_text(self, text: str) -> str:
        """Clean text for better chunking."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()
    
    def chunk_by_sentence(self, text: str) -> List[str]:
        """Chunk text by sentences."""
        text = self._clean_text(text)
        
        # Enhanced sentence splitting pattern
        sentence_pattern = r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=[.!?])\s+'
        sentences = re.split(sentence_pattern, text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed token limit
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            if self._count_tokens(potential_chunk) > self.config.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk = potential_chunk
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def chunk_recursively(self, text: str) -> List[str]:
        """Chunk text using recursive splitting."""
        text = self._clean_text(text)
        
        def split_text_recursive(text: str, separators: List[str]) -> List[str]:
            if not separators:
                # If no separators left, split by token count
                return self._split_by_tokens(text)
            
            separator = separators[0]
            remaining_separators = separators[1:]
            
            if separator not in text:
                return split_text_recursive(text, remaining_separators)
            
            splits = text.split(separator)
            chunks = []
            current_chunk = ""
            
            for split in splits:
                potential_chunk = current_chunk + separator + split if current_chunk else split
                
                if self._count_tokens(potential_chunk) <= self.config.chunk_size:
                    current_chunk = potential_chunk
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    
                    # If this split is still too large, recursively split it
                    if self._count_tokens(split) > self.config.chunk_size:
                        sub_chunks = split_text_recursive(split, remaining_separators)
                        chunks.extend(sub_chunks)
                        current_chunk = ""
                    else:
                        current_chunk = split
            
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            return chunks
        
        chunks = split_text_recursive(text, self.config.separators)
        
        # Add overlap if configured
        if self.config.chunk_overlap > 0:
            chunks = self._add_overlap(chunks)
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def _split_by_tokens(self, text: str) -> List[str]:
        """Split text by token count when no separators work."""
        tokens = self.encoding.encode(text)
        chunks = []
        
        for i in range(0, len(tokens), self.config.chunk_size):
            chunk_tokens = tokens[i:i + self.config.chunk_size]
            chunk_text = self.encoding.decode(chunk_tokens)
            chunks.append(chunk_text)
        
        return chunks
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlap between chunks."""
        if len(chunks) <= 1:
            return chunks
        
        overlapped_chunks = [chunks[0]]
        
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i-1]
            current_chunk = chunks[i]
            
            # Get overlap from previous chunk
            prev_tokens = self.encoding.encode(prev_chunk)
            if len(prev_tokens) > self.config.chunk_overlap:
                overlap_tokens = prev_tokens[-self.config.chunk_overlap:]
                overlap_text = self.encoding.decode(overlap_tokens)
                overlapped_chunk = overlap_text + " " + current_chunk
            else:
                overlapped_chunk = current_chunk
            
            overlapped_chunks.append(overlapped_chunk)
        
        return overlapped_chunks
    
    def chunk_fixed_size(self, text: str) -> List[str]:
        """Chunk text into fixed-size pieces."""
        text = self._clean_text(text)
        tokens = self.encoding.encode(text)
        chunks = []
        
        for i in range(0, len(tokens), self.config.chunk_size - self.config.chunk_overlap):
            end_idx = min(i + self.config.chunk_size, len(tokens))
            chunk_tokens = tokens[i:end_idx]
            chunk_text = self.encoding.decode(chunk_tokens)
            chunks.append(chunk_text)
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def chunk_text(self, text: str) -> List[str]:
        """Chunk text based on configured strategy."""
        if not text or not text.strip():
            return []
        
        if self.config.strategy == ChunkingStrategy.SENTENCE:
            return self.chunk_by_sentence(text)
        elif self.config.strategy == ChunkingStrategy.RECURSIVE:
            return self.chunk_recursively(text)
        elif self.config.strategy == ChunkingStrategy.FIXED_SIZE:
            return self.chunk_fixed_size(text)
        else:
            raise ValueError(f"Unknown chunking strategy: {self.config.strategy}")

class EmbeddingGenerator:
    """Handles embedding generation with OpenAI API."""
    
    def __init__(self, config: EmbeddingConfig):
        self.config = config
        self.client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.encoding.encode(text))
    
    def _validate_text(self, text: str) -> str:
        """Validate and truncate text if necessary."""
        token_count = self._count_tokens(text)
        if token_count > self.config.max_tokens:
            logger.warning(f"Text truncated from {token_count} to {self.config.max_tokens} tokens")
            tokens = self.encoding.encode(text)[:self.config.max_tokens]
            return self.encoding.decode(tokens)
        return text
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for a list of texts."""
        embeddings = []
        
        # Process in batches
        for i in range(0, len(texts), self.config.batch_size):
            batch = texts[i:i + self.config.batch_size]
            validated_batch = [self._validate_text(text) for text in batch if text.strip()]
            
            if not validated_batch:
                continue
            
            try:
                response = self.client.embeddings.create(
                    input=validated_batch,
                    model=self.config.model
                )
                batch_embeddings = [item.embedding for item in response.data]
                embeddings.extend(batch_embeddings)
                
                logger.info(f"Generated embeddings for batch {i//self.config.batch_size + 1}")
                
                # Add small delay to respect rate limits
                time.sleep(0.1)
                
            except Exception as e:
                logger.error(f"Error generating embeddings for batch {i}: {e}")
                raise
        
        return embeddings

class ChromaDBManager:
    """Manages ChromaDB operations."""
    
    def __init__(self, db_path: str):
        self.db_path = db_path
        self.client = chromadb.PersistentClient(path=db_path)
    
    def get_or_create_collection(self, name: str) -> chromadb.Collection:
        """Get or create a ChromaDB collection."""
        try:
            return self.client.get_or_create_collection(name=name)
        except Exception as e:
            logger.error(f"Error creating collection {name}: {e}")
            raise
    
    def add_documents(
        self,
        collection_name: str,
        documents: List[str],
        embeddings: List[List[float]],
        metadatas: List[Dict[str, Any]],
        ids: Optional[List[str]] = None
    ) -> None:
        """Add documents to a ChromaDB collection."""
        if not documents:
            return
        
        collection = self.get_or_create_collection(collection_name)
        
        if ids is None:
            timestamp = int(time.time())
            ids = [f"doc_{timestamp}_{i}" for i in range(len(documents))]
        
        try:
            collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas
            )
            logger.info(f"Added {len(documents)} documents to collection '{collection_name}'")
        except Exception as e:
            logger.error(f"Error adding documents to collection: {e}")
            raise

class RAGDataLoader:
    """Main class for loading data into RAG system."""
    
    def __init__(
        self,
        data_dir: str,
        db_path: str,
        chunking_config: Optional[ChunkingConfig] = None,
        embedding_config: Optional[EmbeddingConfig] = None
    ):
        self.data_dir = Path(data_dir)
        self.processor = DocumentProcessor()
        self.chunker = TextChunker(chunking_config or ChunkingConfig())
        self.embedding_generator = EmbeddingGenerator(embedding_config or EmbeddingConfig())
        self.db_manager = ChromaDBManager(db_path)
        
        # Validate OpenAI API key
        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("OPENAI_API_KEY environment variable is required")
        
        # Ensure data directory exists
        self.data_dir.mkdir(exist_ok=True)
    
    def load_single_file(
        self,
        file_path: Union[str, Path],
        collection_name: str = "user_documents"
    ) -> Dict[str, Any]:
        """Load a single file into the RAG system."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Processing file: {file_path}")
        
        try:
            # Extract text
            text = self.processor.extract_text(file_path)
            logger.info(f"Extracted {len(text)} characters from {file_path}")
            
            if not text.strip():
                raise ValueError(f"No text content extracted from {file_path}")
            
            # Chunk text
            chunks = self.chunker.chunk_text(text)
            logger.info(f"Created {len(chunks)} chunks")
            
            if not chunks:
                raise ValueError(f"No chunks created from {file_path}")
            
            # Generate embeddings
            embeddings = self.embedding_generator.generate_embeddings(chunks)
            
            # Prepare metadata
            metadatas = []
            for i, chunk in enumerate(chunks):
                metadata = {
                    "source": str(file_path),
                    "filename": file_path.name,
                    "file_type": file_path.suffix.lower(),
                    "chunk_index": i,
                    "chunk_size": len(chunk),
                    "tokens": self.chunker._count_tokens(chunk),
                    "created_at": time.time()
                }
                metadatas.append(metadata)
            
            # Generate unique IDs
            timestamp = int(time.time())
            ids = [f"{file_path.stem}_{timestamp}_{i}" for i in range(len(chunks))]
            
            # Store in ChromaDB
            self.db_manager.add_documents(
                collection_name=collection_name,
                documents=chunks,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
            
            return {
                "status": "success",
                "file": str(file_path),
                "chunks": len(chunks),
                "collection": collection_name
            }
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            raise
    
    def load_directory(
        self,
        directory_path: Union[str, Path],
        collection_name: str = "user_documents",
        file_extensions: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """Load all supported files from a directory."""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        if file_extensions is None:
            file_extensions = [".txt", ".pdf", ".docx"]
        
        files = []
        for ext in file_extensions:
            files.extend(directory_path.glob(f"*{ext}"))
            files.extend(directory_path.glob(f"**/*{ext}"))  # Include subdirectories
        
        if not files:
            logger.warning(f"No supported files found in {directory_path}")
            return []
        
        logger.info(f"Found {len(files)} files to process")
        
        results = []
        for file_path in files:
            try:
                result = self.load_single_file(file_path, collection_name)
                results.append(result)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                results.append({
                    "status": "error",
                    "file": str(file_path),
                    "error": str(e)
                })
                continue
        
        logger.info(f"Completed processing directory: {directory_path}")
        return results

# Utility functions for compatibility
def get_chroma_client():
    """Get ChromaDB client instance."""
    data_dir_path = os.getenv("DATA_DIR", "data")
    return chromadb.PersistentClient(path=f"{data_dir_path}/chroma_db")

def get_embedding(text: str, model: str = "text-embedding-3-small") -> List[float]:
    """
    Get embedding for a single text using OpenAI.
    Maintains compatibility with existing code.
    """
    try:
        client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        response = client.embeddings.create(
            input=[text],
            model=model
        )
        return response.data[0].embedding
    except Exception as e:
        logger.error(f"Error generating embedding: {e}")
        raise

def main():
    """Main function for command-line usage."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Load documents into RAG system")
    parser.add_argument("--file", "-f", type=str, help="Single file to process")
    parser.add_argument("--directory", "-d", type=str, help="Directory to process")
    parser.add_argument("--collection", "-c", type=str, default="user_documents", help="Collection name")
    parser.add_argument("--data-dir", type=str, default="data", help="Data directory")
    parser.add_argument("--chunk-size", type=int, default=1000, help="Chunk size")
    parser.add_argument("--chunk-overlap", type=int, default=200, help="Chunk overlap")
    parser.add_argument("--strategy", type=str, default="recursive", 
                       choices=["sentence", "recursive", "fixed_size"], help="Chunking strategy")
    
    args = parser.parse_args()
    
    # Configure chunking
    chunking_config = ChunkingConfig(
        strategy=ChunkingStrategy(args.strategy),
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap
    )
    
    # Initialize loader
    db_path = os.path.join(args.data_dir, "chroma_db")
    loader = RAGDataLoader(
        data_dir=args.data_dir,
        db_path=db_path,
        chunking_config=chunking_config
    )
    
    # Process files
    if args.file:
        result = loader.load_single_file(args.file, args.collection)
        print(f"Result: {result}")
    elif args.directory:
        results = loader.load_directory(args.directory, args.collection)
        for result in results:
            print(f"Result: {result}")
    else:
        logger.error("Either --file or --directory must be specified")

if __name__ == "__main__":
    main()